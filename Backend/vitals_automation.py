import re
import sys
import os
from docx import Document
from datetime import datetime
from docx.shared import Pt
from docx.oxml.ns import qn

import spacy
import medspacy

# ==============================
# LOAD MEDSPACY + MED7
# ==============================
try:
    nlp = medspacy.load()
except Exception:
    nlp = spacy.blank("en")
    nlp.add_pipe("sentencizer")

try:
    med7_nlp = spacy.load("en_core_med7_lg")
except Exception:
    med7_nlp = None


# ==============================
# GLOBAL PATTERNS
# ==============================
DEGREE_PATTERN = (
    r'MD|DO|DC|PT|OT|QME|L\.Ac\.|FAAOS|PA-C|PA|DNP|DPM|FNP-BC|PsyD|'
    r'F\.A\.C\.S|FICS|FACOG|RN|Ph\.?\s?D|NP-C|NP|L\.V\.N|RPT|DPT|'
    r'CNM|RNP|FNP|APRN|FACP|FAAP|PΑ'
)

NAME_TOKEN_REGEX = re.compile(
    r"^[A-Z][a-zA-Z]+(?:[-'][A-Z][a-zA-Z]+|[-'][a-zA-Z]+)*\.?$"
)

DATE_RECORD_REGEX = re.compile(
    r'(?=(?:^|\n)\s*\d{1,2}/\d{1,2}/(?:\d{2,4})?\.)',
    re.MULTILINE
)

STRENGTH_UNIT_PATTERN = (
    r'(?:mg|mcg|g|gm|kg|ml|mL|l|L|iu|IU|units|Unit|meq|mEq|mmol|'
    r'mg/ml|mg/mL|mcg/ml|mcg/mL|g/ml|g/mL|%)'
)

STRENGTH_PATTERN = rf'\b\d+(?:\.\d+)?(?:/\d+(?:\.\d+)?)?\s*{STRENGTH_UNIT_PATTERN}\b'


# ==============================
# 1. READ FULL DOCX TEXT
# ==============================
def read_docx_text(file_path):
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return ""

    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_parts.append(cell_text)
            if row_parts:
                parts.append(" ".join(row_parts))

    return "\n".join(parts)


# ==============================
# 2. EXTRACT ONLY MEDICAL RECORDS SECTION
# ==============================
def get_medical_records_section(full_text):
    match = re.search(r'MEDICAL RECORDS:\s*(.*)', full_text, re.IGNORECASE | re.DOTALL)
    if match:
        return match.group(1)
    return full_text


# ==============================
# 3. SPLIT TEXT INTO RECORDS
# ==============================
def split_records_from_text(text):
    text = text.replace("\r", "\n")

    starts = [m.start() for m in DATE_RECORD_REGEX.finditer(text)]
    records = []

    if not starts:
        return records

    for i, start in enumerate(starts):
        end = starts[i + 1] if i + 1 < len(starts) else len(text)

        record_text = text[start:end].strip()

        # important: keep enough text before record for names on prior line
        before_context = text[max(0, start - 300):start].strip()

        record_text_clean = " ".join(record_text.split()).strip()
        before_context_clean = " ".join(before_context.split()).strip()

        if re.match(r'^\d{1,2}/\d{1,2}/(?:\d{2,4})?\.', record_text_clean):
            records.append({
                "record_text": record_text_clean,
                "before_context": before_context_clean
            })

    return records


# ==============================
# 4. DATE EXTRACTION
# ==============================
def extract_date(record):
    m = re.match(r'^\s*(\d{1,2}/\d{1,2}/(?:\d{2,4})?)\.', record)
    if not m:
        return ""

    date = m.group(1)
    if re.match(r'^\d{1,2}/\d{1,2}/$', date):
        return date + "00"
    return date


# ==============================
# 5. PROVIDER EXTRACTION
# ==============================
def clean_provider(provider):
    provider = provider.strip()
    provider = re.sub(r'\s*\([^)]*\)', '', provider).strip()
    provider = re.sub(r'\s+,', ',', provider)
    provider = re.sub(r'\s+', ' ', provider).strip()
    return provider


def extract_provider(record_text, before_context=""):
    record_text = clean_spaces(record_text)
    before_context = clean_spaces(before_context)

    # 1. explicit unknown provider
    if re.search(r'\bUnknown\s+Provider\b', record_text, re.IGNORECASE):
        return "Unknown Provider"

    full_text = f"{before_context}\n{record_text}".strip()

    # provider name + degree
    provider_pattern = re.compile(
        rf"""
        \b
        ([A-Z][a-zA-Z]+(?:[-'][A-Za-z]+)?(?:\s+[A-Z][a-zA-Z]+(?:[-'][A-Za-z]+)?)+)
        ,?\s*
        ({DEGREE_PATTERN})
        \b
        """,
        re.VERBOSE | re.IGNORECASE
    )

    # common non-provider words to reject
    bad_provider_words = {
        "medical", "center", "clinic", "foundation", "hospital",
        "care", "group", "corporation", "services", "report",
        "note", "records", "industrial", "permanente"
    }

    def valid_provider(name, degree):
        provider = f"{name.strip()}, {degree.strip().upper()}"
        provider = re.sub(r'\s*\([^)]*\)', '', provider).strip()
        provider = re.sub(r'\s+,', ',', provider)
        provider = re.sub(r'\s+', ' ', provider).strip()

        low = provider.lower()
        if any(word in low for word in bad_provider_words):
            return ""
        return provider

    # 2. FIRST: provider in before_context line above the record
    # Example:
    # Karita Goulbourne, MD
    # 03/22/12. ED Provider Note...
    context_matches = provider_pattern.findall(before_context)
    for name, degree in reversed(context_matches):
        provider = valid_provider(name, degree)
        if provider:
            return provider

    # 3. SECOND: provider inside the current record
    # Example:
    # 05/20/20. UNIVERSAL INDUSTRIAL CARE. Oscar Tuazon, MD...
    record_matches = provider_pattern.findall(record_text)
    for name, degree in record_matches:
        provider = valid_provider(name, degree)
        if provider:
            return provider

    # 4. THIRD: search combined text as fallback
    all_matches = provider_pattern.findall(full_text)
    for name, degree in all_matches:
        provider = valid_provider(name, degree)
        if provider:
            return provider

    return "Unknown Provider"


# ==============================
# 6. BASIC FIELD EXTRACTION
# ==============================
def extract_bp(text):
    m = re.search(r'\bBP[:\s]*([0-9]{2,3}/[0-9]{2,3})\b', text, re.IGNORECASE)
    if m:
        return m.group(1)

    matches = re.findall(r'\b([0-9]{2,3}/[0-9]{2,3})\b', text)
    for val in matches:
        try:
            s, d = map(int, val.split('/'))
            if 80 <= s <= 250 and 50 <= d <= 150:
                return val
        except Exception:
            pass
    return ""


def extract_weight(text):
    kg_patterns = [
        r'\bWt[:\s]*([0-9]{2,3}(?:\.\d+)?)\s*(kg|kgs)\b',
        r'\bWeight[:\s]*([0-9]{2,3}(?:\.\d+)?)\s*(kg|kgs)\b',
        r'\b([0-9]{2,3}(?:\.\d+)?)\s*(kg|kgs)\b'
    ]

    for pat in kg_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            kg = float(m.group(1))
            lbs = kg * 2.20462
            return str(round(lbs, 2))

    lbs_patterns = [
        r'\bWt[:\s]+([0-9]{2,3}(?:\.\d+)?)\b',
        r'\bWt[:\s]+([0-9]{2,3}(?:\.\d+)?)(?:\s?lbs?)\b',
        r'\bWt\s*(?:lbs?|lb)\s*[:\-]?\s*([0-9]{2,3}(?:\.\d+)?)\b',
        r'\bWeight[:\s]+([0-9]{2,3}(?:\.\d+)?)\b',
        r'\bWeight[:\s]+([0-9]{2,3}(?:\.\d+)?)(?:\s?lbs?)\b',
        r'\bWeight\s*(?:lbs?|lb)\s*[:\-]?\s*([0-9]{2,3}(?:\.\d+)?)\b',
    ]

    for pat in lbs_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(1)

    return ""


def extract_spo2(text):
    patterns = [
        r'\bSpO2[:\s]*([0-9]{1,3})\b',
        r'\bSPO2[:\s]*([0-9]{1,3})\b',
        r'\bSpO₂[:\s]*([0-9]{1,3})\b',
        r'\bO2\s*Sat[:\s]*([0-9]{1,3})\b',
        r'\bO2\s*SAT[:\s]*([0-9]{1,3})\b',
        r'\bSAT[:\s]*([0-9]{1,3})\b',
        r'\bSat[:\s]*([0-9]{1,3})\b',
        r'\bOxygen\s*Sat(?:uration)?[:\s]*([0-9]{1,3})\b'
    ]

    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            val = int(m.group(1))
            if 50 <= val <= 100:
                return str(val)

    return ""


def extract_sugar(text):
    patterns = [
        r'\bGlucose[:\s]*([0-9]+(?:\.\d+)?)\b',
        r'\bSugar[:\s]*([0-9]+(?:\.\d+)?)\b'
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(1)
    return ""


def extract_a1c(text):
    patterns = [
        r'\bHBA1C[:\s]*([0-9]+(?:\.\d+)?)\b',
        r'\bHbA1C[:\s]*([0-9]+(?:\.\d+)?)\b',
        r'\bHgA1C[:\s]*([0-9]+(?:\.\d+)?)\b',
        r'\bA1C[:\s]*([0-9]+(?:\.\d+)?)\b'
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(1)
    return ""


# ==============================
# 7. MEDICATION EXTRACTION
# ==============================
DOSAGE_FORMS = [
    "eye drops", "ear drops", "nasal drops", "drops", "drop",
    "ophthalmic solution", "ophthalmic ointment",
    "ointment", "cream", "gel", "lotion",
    "tablet", "tablets", "tab", "tabs",
    "capsule", "capsules", "cap", "caps",
    "syrup", "suspension", "solution", "oral solution",
    "injection", "inj", "iv", "intravenous", "iv fluid", "flush",
    "spray", "inhaler", "patch", "suppository", "nebulizer",
    "powder", "elixir", "lozenge", "kit"
]

ACTION_WORDS = [
    "prescribed", "given", "started", "administered", "continue",
    "continued", "refilled", "take", "use", "apply", "instill",
    "inject", "treated with", "discharged with", "start"
]


def clean_spaces(text):
    return " ".join(text.split()).strip()


def split_med_sentences(text):
    try:
        doc = nlp(text)
        sents = [clean_spaces(sent.text) for sent in doc.sents if clean_spaces(sent.text)]
        if sents:
            return sents
    except Exception:
        pass

    parts = re.split(r'[\n;]+|(?<=[.])\s+', text)
    return [clean_spaces(p) for p in parts if clean_spaces(p)]


def get_med_plan_sections(text):
    found = []

    patterns = [
        (
            "current medications",
            r'Current\s+Medications?:\s*(.*?)(?=(?:Assessment/Plan:|Assessment:|Plan:|Discussion:|Subjective:|Objective:|HPI:|Diagnosis:|Diagnoses:|Chief Complaint:|History of Present Illness:|Past Medical History:|Vital Signs:|Vitals:|Work Status:|$))'
        ),
        (
            "medications",
            r'Medications?:\s*(.*?)(?=(?:Assessment/Plan:|Assessment:|Plan:|Discussion:|Subjective:|Objective:|HPI:|Diagnosis:|Diagnoses:|Chief Complaint:|History of Present Illness:|Past Medical History:|Vital Signs:|Vitals:|Work Status:|$))'
        ),
        (
            "plan",
            r'Plan:\s*(.*?)(?=(?:Assessment/Plan:|Assessment:|Discussion:|Subjective:|Objective:|HPI:|Diagnosis:|Diagnoses:|Chief Complaint:|History of Present Illness:|Past Medical History:|Vital Signs:|Vitals:|Work Status:|$))'
        ),
        (
            "discussion",
            r'Discussion:\s*(.*?)(?=(?:Assessment/Plan:|Assessment:|Plan:|Subjective:|Objective:|HPI:|Diagnosis:|Diagnoses:|Chief Complaint:|History of Present Illness:|Past Medical History:|Vital Signs:|Vitals:|Work Status:|$))'
        ),
    ]

    for label, pat in patterns:
        matches = re.findall(pat, text, re.IGNORECASE | re.DOTALL)
        for m in matches:
            cleaned = clean_spaces(m)
            if cleaned:
                found.append((label, cleaned))

    return found


def normalize_med_name_case(name):
    if not name:
        return ""
    words = []
    for w in name.split():
        if re.search(r'\d', w):
            if w.lower().startswith("vitamin"):
                words.append("Vitamin")
            else:
                words.append(w)
        else:
            if len(w) <= 3 and w.isupper():
                words.append(w)
            else:
                words.append(w[0].upper() + w[1:] if w else w)
    return " ".join(words)


def clean_candidate_name(name):
    name = clean_spaces(name)

    name = re.sub(r'^(?:an|a|the)\s+', '', name, flags=re.IGNORECASE)
    name = re.sub(
        r'^(?:' + "|".join(re.escape(x) for x in ACTION_WORDS) + r')\s+',
        '',
        name,
        flags=re.IGNORECASE
    )

    dosage_forms_regex = "|".join(re.escape(x) for x in DOSAGE_FORMS)
    name = re.sub(rf'\b(?:{dosage_forms_regex})\b', '', name, flags=re.IGNORECASE)

    name = re.sub(
        r'\b(?:po|by mouth|daily|once daily|twice daily|three times daily|qday|qd|bid|tid|qid|qhs|prn|on hold|hold|held|oral)\b',
        '',
        name,
        flags=re.IGNORECASE
    )

    name = re.sub(
        r'\b(?:1 tablet|2 tablets|1 capsule|2 capsules|1 cap|2 caps|1 tab|2 tabs)\b',
        '',
        name,
        flags=re.IGNORECASE
    )

    name = re.sub(
        r'\b(?:was|were|is|are|for|and|or|to|by|with|patient|plan|discussion|assessment)\b$',
        '',
        name,
        flags=re.IGNORECASE
    )

    name = re.sub(r'\s+', ' ', name).strip(" ,.;:-")
    return normalize_med_name_case(name)


def looks_like_real_medication(candidate, source_section=""):
    candidate = clean_spaces(candidate)
    if not candidate:
        return False

    words = re.findall(r'[A-Za-z][A-Za-z0-9\-]*', candidate)
    if not words:
        return False

    if len(words) > 4:
        return False

    if re.fullmatch(STRENGTH_PATTERN, candidate, re.IGNORECASE):
        return False

    if source_section.lower() in ["current medications", "medications"]:
        return True

    has_strength = bool(re.search(STRENGTH_PATTERN, candidate, re.IGNORECASE))
    has_dosage_form = bool(re.search(
        r'\b(tablet|tab|capsule|cap|drops|drop|solution|ointment|cream|gel|patch|spray|inhaler|injection|iv|syrup|suspension|saline)\b',
        candidate,
        re.IGNORECASE
    ))

    return has_strength or has_dosage_form


def extract_name_and_strength(text, section_label=""):
    text = clean_spaces(text)

    strength_match = re.search(STRENGTH_PATTERN, text, re.IGNORECASE)
    strength = strength_match.group(0).strip() if strength_match else ""

    name = re.sub(STRENGTH_PATTERN, '', text, flags=re.IGNORECASE)
    name = clean_candidate_name(name)

    if not name:
        return ""

    result = f"{name} {strength}".strip() if strength else name

    if not looks_like_real_medication(result, section_label):
        return ""

    return result


def dedupe_keep_order(items):
    out = []
    seen = set()
    for item in items:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            out.append(item)
    return out


def extract_candidate_meds_med7(section_text, section_label=""):
    candidates = []

    if med7_nlp is None:
        return candidates

    for sent in split_med_sentences(section_text):
        try:
            doc = med7_nlp(sent)
        except Exception:
            continue

        for ent in doc.ents:
            if ent.label_.upper().strip() == "DRUG":
                med_name = extract_name_and_strength(ent.text, section_label)
                if med_name:
                    candidates.append(med_name)

    return candidates


def extract_candidate_meds_regex(section_text, section_label=""):
    candidates = []

    sentences = split_med_sentences(section_text)

    for sent in sentences:
        # 1. name + strength
        for m in re.finditer(
            rf'\b([A-Za-z][A-Za-z0-9\-]*(?:\s+[A-Za-z][A-Za-z0-9\-]*){{0,2}})\s+({STRENGTH_PATTERN})',
            sent,
            re.IGNORECASE
        ):
            med = extract_name_and_strength(m.group(0), section_label)
            if med:
                candidates.append(med)

        # 2. medication sections: split all pieces
        if section_label.lower() in ["current medications", "medications"]:
            parts = re.split(r'[;,\n]+', sent)

            for part in parts:
                part = clean_spaces(part)
                if not part:
                    continue

                med = extract_name_and_strength(part, section_label)
                if med:
                    candidates.append(med)

        # 3. plan / discussion: action-based extraction
        else:
            for m in re.finditer(
                rf'(?:{"|".join([re.escape(x) for x in ACTION_WORDS])})\s+'
                rf'([A-Za-z][A-Za-z0-9\-]*(?:\s+[A-Za-z0-9\-]+){{0,4}}(?:\s+{STRENGTH_PATTERN})?)',
                sent,
                re.IGNORECASE
            ):
                med = extract_name_and_strength(m.group(1), section_label)
                if med:
                    candidates.append(med)

            for m in re.finditer(
                r'\b(normal\s+saline|saline|sodium\s+chloride(?:\s+0\.9%)?)\b',
                sent,
                re.IGNORECASE
            ):
                med = extract_name_and_strength(m.group(1), section_label)
                if med:
                    candidates.append(med)

    return candidates


def extract_medication(text):
    sections = get_med_plan_sections(text)
    if not sections:
        return ""

    results = []

    for label, section_text in sections:
        med7_candidates = extract_candidate_meds_med7(section_text, label)
        regex_candidates = extract_candidate_meds_regex(section_text, label)

        for med_name in med7_candidates + regex_candidates:
            med_name = clean_spaces(med_name)
            if med_name and looks_like_real_medication(med_name, label):
                results.append(med_name)

    return ", ".join(dedupe_keep_order(results))


# ==============================
# 8. FILTER NON-MEDICAL RECORDS
# ==============================
def is_medical_record(record_text, before_context=""):
    rec = record_text.lower()

    excluded = [
        "attestation",
        "cover letter",
        "declaration",
        "reviewed but not summarized"
    ]
    for x in excluded:
        if x in rec:
            return False

    included = [
        "progress note",
        "ed provider note",
        "ed note",
        "provider note",
        "lab report",
        "subjective",
        "chief complaint",
        "present complaints",
        "patient complaints",
        "patient complaint",
        "hpi",
        "vitals",
        "vital signs",
        "vital",
        "diagnosis",
        "diagnoses",
        "plan:",
        "medications:",
        "current medications:",
        "discussion:"
    ]

    if any(x in rec for x in included):
        return True

    if (
        extract_bp(record_text)
        or extract_weight(record_text)
        or extract_spo2(record_text)
        or extract_sugar(record_text)
        or extract_a1c(record_text)
        or extract_medication(record_text)
        or extract_provider(record_text, before_context) != "Unknown Provider"
    ):
        return True

    return False


# ==============================
# 9. DATE SORT
# ==============================
def normalize_date(date_str):
    if not date_str:
        return None

    cleaned = date_str.strip().rstrip(".")
    for fmt in ("%m/%d/%y", "%m/%d/%Y"):
        try:
            return datetime.strptime(cleaned, fmt)
        except Exception:
            pass
    return None


# ==============================
# 10. PROCESS RECORDS
# ==============================
def process_records(records):
    output = []
    processed = []

for i, record in enumerate(records, 1):
        print(f"=== CHECKING RECORD {i} ===")
        print(record[:300])

    for item in records:
        record_text = item["record_text"]
        before_context = item.get("before_context", "")

        if not is_medical_record(record_text, before_context):
            continue

        row = {
            "Date": extract_date(record_text),
            "Provider": extract_provider(record_text, before_context),
            "Bp mmHg": extract_bp(record_text),
            "Wt Lbs": extract_weight(record_text),
            "Spo2": extract_spo2(record_text),
            "Sugar": extract_sugar(record_text),
            "A1c": extract_a1c(record_text),
            "Medication": extract_medication(record_text),
        }

        if any([
            row["Bp mmHg"],
            row["Wt Lbs"],
            row["Spo2"],
            row["Sugar"],
            row["A1c"],
            row["Medication"]
        ]):
            output.append(row)

    output.sort(key=lambda x: normalize_date(x["Date"]) or datetime.max)
    return output
return processed

# ==============================
# 11. WORD TABLE
# ==============================
def set_times_new_roman(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)


def create_word_table(data, output_file):
    doc = Document()

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"

    headers = ["Date", "Provider", "Bp mmHg", "Wt Lbs", "Spo2", "Sugar", "A1c", "Medication"]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_times_new_roman(cell)

    for row in data:
        cells = table.add_row().cells
        cells[0].text = row["Date"]
        cells[1].text = row["Provider"]
        cells[2].text = row["Bp mmHg"]
        cells[3].text = row["Wt Lbs"]
        cells[4].text = row["Spo2"]
        cells[5].text = row["Sugar"]
        cells[6].text = row["A1c"]
        cells[7].text = row["Medication"]

        for cell in cells:
            set_times_new_roman(cell)

    doc.save(output_file)


# ==============================
# 12. MAIN
# ==============================
def main():
    if len(sys.argv) < 2:
        print("Usage: python vitals_automation.py input.docx")
        return

    input_file = sys.argv[1].strip().strip('"')

    if not os.path.exists(input_file):
        print("❌ Input file not found")
        return

    base_name = os.path.splitext(os.path.basename(input_file))[0]
    output_folder = "outputs"
    os.makedirs(output_folder, exist_ok=True)
    output_file = os.path.join(output_folder, f"{base_name}_output.docx")

    full_text = read_docx_text(input_file)
    if not full_text:
        print("❌ No text found in file.")
        return

    medical_text = get_medical_records_section(full_text)
    records = split_records_from_text(medical_text)

    print(f"Total records found: {len(records)}")

    processed = process_records(records)
    print(f"Medical rows extracted: {len(processed)}")

    if not processed:
        print("❌ No medical records extracted.")
        return

    create_word_table(processed, output_file)
    print(f"✅ Output saved in outputs folder: {output_file}")


def process_file(input_path, output_path):
    print("=== PROCESS FILE START ===")
    print("INPUT:", input_path)
    print("OUTPUT:", output_path)

    text = read_docx(input_path)   # or your read_file logic
    print("=== TEXT LENGTH ===", len(text) if text else 0)

    if not text or not text.strip():
        raise Exception("No text extracted from file")

    records = split_records(text)
    print("=== RECORD COUNT ===", len(records))

    processed = process_records(records)
    print("=== PROCESSED COUNT ===", len(processed))

    if not processed:
        raise Exception("No medical records extracted")

    save_output(processed, output_path)
    print("=== PROCESS FILE END ===")

if __name__ == "__main__":
    main()